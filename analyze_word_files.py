import docx
import os
from pathlib import Path

def analyze_word_file(file_path):
    """Analyze a Word document and extract its structure and content"""
    try:
        doc = docx.Document(file_path)
        print(f"\n{'='*80}")
        print(f"ANALYZING: {file_path}")
        print(f"{'='*80}")
        
        print(f"Document structure:")
        print(f"- Paragraphs: {len(doc.paragraphs)}")
        print(f"- Tables: {len(doc.tables)}")
        print(f"- Sections: {len(doc.sections)}")
        
        # Analyze paragraphs
        print(f"\nPARAGRAPH ANALYSIS:")
        print(f"{'-'*40}")
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip():  # Only show non-empty paragraphs
                print(f"Paragraph {i+1}: {para.text[:200]}{'...' if len(para.text) > 200 else ''}")
                if i >= 20:  # Limit to first 20 paragraphs
                    print(f"... (showing first 20 paragraphs)")
                    break
        
        # Analyze tables
        print(f"\nTABLE ANALYSIS:")
        print(f"{'-'*40}")
        for i, table in enumerate(doc.tables):
            print(f"\nTable {i+1}:")
            print(f"- Rows: {len(table.rows)}")
            print(f"- Columns: {len(table.columns) if table.rows else 0}")
            
            # Show first few rows of each table
            for j, row in enumerate(table.rows[:5]):  # Show first 5 rows
                row_data = [cell.text.strip() for cell in row.cells]
                print(f"  Row {j+1}: {row_data}")
            if len(table.rows) > 5:
                print(f"  ... (showing first 5 rows)")
        
        # Look for specific patterns
        print(f"\nCONTENT PATTERNS:")
        print(f"{'-'*40}")
        
        # Search for time patterns
        time_patterns = []
        for para in doc.paragraphs:
            text = para.text.lower()
            if any(time_word in text for time_word in ['am', 'pm', ':', 'time', 'schedule']):
                time_patterns.append(para.text)
        
        if time_patterns:
            print(f"Time-related content found:")
            for pattern in time_patterns[:5]:
                print(f"  - {pattern[:100]}{'...' if len(pattern) > 100 else ''}")
        
        # Search for session types
        session_types = []
        for para in doc.paragraphs:
            text = para.text.lower()
            if any(session_word in text for session_word in ['lecture', 'session', 'presentation', 'talk', 'workshop', 'panel']):
                session_types.append(para.text)
        
        if session_types:
            print(f"Session-related content found:")
            for session in session_types[:5]:
                print(f"  - {session[:100]}{'...' if len(session) > 100 else ''}")
        
        # Search for speaker/author patterns
        speaker_patterns = []
        for para in doc.paragraphs:
            text = para.text
            if any(speaker_word in text.lower() for speaker_word in ['dr.', 'prof.', 'mr.', 'ms.', 'presenter', 'speaker']):
                speaker_patterns.append(para.text)
        
        if speaker_patterns:
            print(f"Speaker-related content found:")
            for speaker in speaker_patterns[:5]:
                print(f"  - {speaker[:100]}{'...' if len(speaker) > 100 else ''}")
        
        return {
            'paragraphs': len(doc.paragraphs),
            'tables': len(doc.tables),
            'has_time_patterns': len(time_patterns) > 0,
            'has_session_patterns': len(session_types) > 0,
            'has_speaker_patterns': len(speaker_patterns) > 0
        }
        
    except Exception as e:
        print(f"Error analyzing {file_path}: {str(e)}")
        return None

def main():
    """Main function to analyze all Word files"""
    files_dir = Path("files")
    
    if not files_dir.exists():
        print("Files directory not found!")
        return
    
    word_files = list(files_dir.glob("*.docx"))
    
    if not word_files:
        print("No Word files found in the files directory!")
        return
    
    print(f"Found {len(word_files)} Word files to analyze:")
    for file in word_files:
        print(f"  - {file.name}")
    
    results = []
    for file in word_files:
        result = analyze_word_file(file)
        if result:
            results.append((file.name, result))
    
    # Summary
    print(f"\n{'='*80}")
    print(f"SUMMARY ANALYSIS")
    print(f"{'='*80}")
    
    for filename, result in results:
        print(f"\n{filename}:")
        print(f"  - Paragraphs: {result['paragraphs']}")
        print(f"  - Tables: {result['tables']}")
        print(f"  - Has time patterns: {result['has_time_patterns']}")
        print(f"  - Has session patterns: {result['has_session_patterns']}")
        print(f"  - Has speaker patterns: {result['has_speaker_patterns']}")

if __name__ == "__main__":
    main()
