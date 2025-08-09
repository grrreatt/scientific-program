import docx
import os
from pathlib import Path
import re

def extract_table_content(file_path):
    """Extract detailed content from tables in Word documents"""
    try:
        doc = docx.Document(file_path)
        print(f"\n{'='*100}")
        print(f"DETAILED ANALYSIS: {file_path}")
        print(f"{'='*100}")
        
        for i, table in enumerate(doc.tables):
            print(f"\nTABLE {i+1} - COMPLETE CONTENT:")
            print(f"{'='*60}")
            
            # Get table dimensions
            rows = len(table.rows)
            cols = len(table.columns) if table.rows else 0
            print(f"Dimensions: {rows} rows × {cols} columns")
            
            # Extract all content
            table_data = []
            for j, row in enumerate(table.rows):
                row_data = []
                for k, cell in enumerate(row.cells):
                    cell_text = cell.text.strip()
                    row_data.append(cell_text)
                table_data.append(row_data)
                print(f"Row {j+1}: {row_data}")
            
            # Analyze structure
            analyze_table_structure(table_data, file_path, i+1)
            
    except Exception as e:
        print(f"Error analyzing {file_path}: {str(e)}")

def analyze_table_structure(table_data, file_path, table_num):
    """Analyze the structure of a table to understand the program format"""
    print(f"\nSTRUCTURE ANALYSIS for Table {table_num}:")
    print(f"{'='*50}")
    
    if not table_data:
        return
    
    # Look for headers
    headers = table_data[0] if table_data else []
    print(f"Headers: {headers}")
    
    # Identify time patterns
    time_slots = []
    sessions = []
    speakers = []
    halls = []
    
    for i, row in enumerate(table_data):
        for j, cell in enumerate(row):
            cell_lower = cell.lower()
            
            # Time patterns
            if re.search(r'\d{1,2}:\d{2}', cell) or 'am' in cell_lower or 'pm' in cell_lower:
                time_slots.append((i, j, cell))
            
            # Session patterns
            if any(word in cell_lower for word in ['session', 'lecture', 'symposium', 'workshop', 'panel']):
                sessions.append((i, j, cell))
            
            # Speaker patterns
            if any(word in cell_lower for word in ['dr.', 'prof.', 'chairperson', 'moderator']):
                speakers.append((i, j, cell))
            
            # Hall patterns
            if any(word in cell_lower for word in ['hall', 'auditorium', 'room']):
                halls.append((i, j, cell))
    
    print(f"\nTime slots found: {len(time_slots)}")
    for time in time_slots[:5]:
        print(f"  - Row {time[0]+1}, Col {time[1]+1}: {time[2]}")
    
    print(f"\nSessions found: {len(sessions)}")
    for session in sessions[:5]:
        print(f"  - Row {session[0]+1}, Col {session[1]+1}: {session[2]}")
    
    print(f"\nSpeakers found: {len(speakers)}")
    for speaker in speakers[:5]:
        print(f"  - Row {speaker[0]+1}, Col {speaker[1]+1}: {speaker[2]}")
    
    print(f"\nHalls found: {len(halls)}")
    for hall in halls[:5]:
        print(f"  - Row {hall[0]+1}, Col {hall[1]+1}: {hall[2]}")

def extract_sessions_from_table(table_data):
    """Extract structured session data from table"""
    sessions = []
    
    if not table_data or len(table_data) < 2:
        return sessions
    
    # Assume first row is headers
    headers = table_data[0]
    print(f"Table headers: {headers}")
    
    # Process data rows
    for i, row in enumerate(table_data[1:], 1):
        if len(row) >= 2:  # At least time and content columns
            time_col = row[0] if row[0] else ""
            content_col = row[1] if len(row) > 1 and row[1] else ""
            
            # Skip empty rows
            if not time_col.strip() and not content_col.strip():
                continue
            
            session_info = {
                'row': i,
                'time': time_col,
                'content': content_col,
                'additional_cols': row[2:] if len(row) > 2 else []
            }
            sessions.append(session_info)
    
    return sessions

def main():
    """Main function to perform detailed analysis"""
    files_dir = Path("files")
    
    if not files_dir.exists():
        print("Files directory not found!")
        return
    
    word_files = list(files_dir.glob("*.docx"))
    
    if not word_files:
        print("No Word files found in the files directory!")
        return
    
    print(f"Performing detailed analysis of {len(word_files)} Word files...")
    
    all_sessions = []
    
    for file in word_files:
        try:
            doc = docx.Document(file)
            
            for i, table in enumerate(doc.tables):
                # Extract table data
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                
                # Extract sessions from this table
                sessions = extract_sessions_from_table(table_data)
                all_sessions.extend(sessions)
                
                print(f"\nExtracted {len(sessions)} sessions from {file.name} - Table {i+1}")
                
        except Exception as e:
            print(f"Error processing {file}: {str(e)}")
    
    # Summary of all extracted sessions
    print(f"\n{'='*100}")
    print(f"TOTAL SESSIONS EXTRACTED: {len(all_sessions)}")
    print(f"{'='*100}")
    
    # Show sample sessions
    print(f"\nSAMPLE SESSIONS:")
    for i, session in enumerate(all_sessions[:10]):
        print(f"\nSession {i+1}:")
        print(f"  Time: {session['time']}")
        print(f"  Content: {session['content'][:100]}{'...' if len(session['content']) > 100 else ''}")
        if session['additional_cols']:
            print(f"  Additional: {session['additional_cols']}")
    
    # Analyze patterns
    time_patterns = set()
    session_types = set()
    
    for session in all_sessions:
        # Time patterns
        if session['time']:
            time_patterns.add(session['time'])
        
        # Session types
        content_lower = session['content'].lower()
        if 'session' in content_lower:
            session_types.add('session')
        if 'lecture' in content_lower:
            session_types.add('lecture')
        if 'symposium' in content_lower:
            session_types.add('symposium')
        if 'workshop' in content_lower:
            session_types.add('workshop')
        if 'panel' in content_lower:
            session_types.add('panel')
    
    print(f"\nTIME PATTERNS FOUND:")
    for pattern in sorted(time_patterns):
        print(f"  - {pattern}")
    
    print(f"\nSESSION TYPES FOUND:")
    for session_type in sorted(session_types):
        print(f"  - {session_type}")

if __name__ == "__main__":
    main()
